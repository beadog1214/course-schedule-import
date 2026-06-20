#!/usr/bin/env python3
"""
课表图片/PDF/Word → iOS/Android 日历 .ics 文件
用法: python3 schedule2ics.py <文件路径> [--name 姓名] [--start 2026-02-23]
支持: PNG JPG PDF DOCX
"""

import json, os, re, subprocess, sys, datetime, tempfile
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent

# ====== 默认时间槽 ======
DEFAULT_SLOTS = {
    "1-2":   ("08:00", "09:40"),
    "3-4":   ("10:00", "11:40"),
    "5-6":   ("13:30", "15:10"),
    "7-8":   ("15:20", "17:00"),
    "9-10":  ("18:30", "20:00"),
    "9-11":  ("18:30", "21:00"),
    "11-12": ("20:10", "21:40"),
}

DAY_MAP = {"一": 0, "二": 1, "三": 2, "四": 3, "五": 4, "六": 5, "日": 6}
DAY_NAMES = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"]


def extract_text(filepath: str) -> str:
    """根据文件类型提取文字：图片OCR / PDF文本 / PDF转图OCR / Word文本"""
    ext = Path(filepath).suffix.lower()

    if ext in ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif', '.webp'):
        return ocr_plain(filepath)

    if ext == '.pdf':
        # 先试直接提取文本
        text = _pdf_text(filepath)
        if text.strip():
            return text
        # 扫面件PDF，转图片OCR
        print("  PDF 无文本层，转图片 OCR...")
        imgs = _pdf_to_images(filepath)
        if imgs:
            all_text = []
            for i, img in enumerate(imgs):
                print(f"  OCR 第 {i+1}/{len(imgs)} 页...")
                t = ocr_image(img)
                if t.strip():
                    all_text.append(t)
                os.remove(img)
            return "\n".join(all_text)
        return ""

    if ext == '.docx':
        return _docx_text(filepath)

    print(f"  ⚠ 不支持的文件类型: {ext}")
    return ""


def _pdf_text(filepath: str) -> str:
    """pdfplumber 提取 PDF 文本"""
    try:
        import pdfplumber
        lines = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    lines.append(t)
        return "\n".join(lines)
    except ImportError:
        print("  ⚠ pdfplumber 未安装，尝试 pdftotext...")
        return _pdftotext(filepath)


def _pdftotext(filepath: str) -> str:
    """pdftotext 命令行回退"""
    try:
        r = subprocess.run(["pdftotext", "-layout", filepath, "-"], capture_output=True, text=True, timeout=30)
        return r.stdout
    except FileNotFoundError:
        print("  ⚠ pdftotext 也未安装")
        print("    安装: pip install pdfplumber  或  brew install poppler")
        return ""


def _pdf_to_images(filepath: str) -> list:
    """PDF 转 PNG 图片列表（用于扫描件 OCR）"""
    try:
        import fitz  # PyMuPDF
        imgs = []
        doc = fitz.open(filepath)
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=200)
            tmp = f"{tempfile.gettempdir()}/pdf_page_{i}.png"
            pix.save(tmp)
            imgs.append(tmp)
        return imgs
    except ImportError:
        pass
    # pdftoppm 回退
    try:
        prefix = f"{tempfile.gettempdir()}/pdf_page"
        subprocess.run(["pdftoppm", "-png", "-r", "200", filepath, prefix], check=True, capture_output=True)
        return sorted(Path(tempfile.gettempdir()).glob("pdf_page*.png"))
    except:
        return []


def _docx_text(filepath: str) -> str:
    """python-docx 提取 Word 文本"""
    try:
        from docx import Document
        doc = Document(filepath)
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        # 也读表格
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))
        return "\n".join(lines)
    except ImportError:
        print("  ⚠ python-docx 未安装")
        print("    pip install python-docx")
        return ""


def ocr_image(image_path: str) -> str:
    """单张图片 OCR，复用 ocr_plain 逻辑"""
    return ocr_plain(image_path)


def ocr_plain(image_path: str) -> str:
    """tesseract OCR，自动检测中文包并提示安装"""
    # 检查中文语言包
    try:
        langs = subprocess.run(["tesseract", "--list-langs"], capture_output=True, text=True).stdout
    except FileNotFoundError:
        print("  ⚠ tesseract 未安装")
        print("    Mac: brew install tesseract tesseract-lang")
        print("    Linux: apt install tesseract-ocr tesseract-ocr-chi-sim")
        print("    Win: https://github.com/UB-Mannheim/tesseract/wiki 下载安装时勾选中文")
        return ""

    if "chi_sim" not in langs:
        print("  ⚠ tesseract 中文语言包未安装")
        print("    Mac: brew install tesseract-lang")
        print("    Linux: apt install tesseract-ocr-chi-sim")
        print("    Win: 重新运行安装程序，勾选 Chinese Simplified")
        return ""

    try:
        result = subprocess.run(
            ["tesseract", image_path, "stdout", "-l", "chi_sim+eng", "--psm", "6"],
            capture_output=True, text=True, timeout=60
        )
        return result.stdout
    except Exception as e:
        print(f"  ⚠ tesseract OCR 失败: {e}")
        return ""


def detect_semester(text: str) -> tuple:
    """从 OCR 文字中捕捉学年学期，如 '2025-2026学年第二学期' → ('2025-2026', '2')"""
    # 模式: "2025-2026学年第2学期" 或 "2025-2026学年第二学期"
    m = re.search(r'(\d{4})\s*[-~～]\s*(\d{4})\s*学[年期].*?第\s*(\d|[一二三四五六七八九十])\s*学[年期]', text)
    if m:
        y1, y2, sem = m.group(1), m.group(2), m.group(3)
        cn = {"一":"1","二":"2","三":"3","四":"4","五":"5","六":"6","七":"7","八":"8","九":"9","十":"10"}
        sem = cn.get(sem, sem)
        return (f"{y1}-{y2}", sem)
    # 模式: "2025-2026-2"
    m = re.search(r'(\d{4})\s*[-~～]\s*(\d{4})\s*[-~～]\s*(\d)', text)
    if m:
        return (f"{m.group(1)}-{m.group(2)}", m.group(3))
    # 模式: 单独的 "2026年春季学期"
    m = re.search(r'(\d{4}).*?春', text)
    if m:
        return (f"{int(m.group(1))-1}-{m.group(1)}", "2")
    m = re.search(r'(\d{4}).*?秋', text)
    if m:
        return (f"{m.group(1)}-{int(m.group(1))+1}", "1")
    return ("", "")


def detect_slots(text: str) -> dict:
    """从 OCR 文字中自动捕捉节次时间，如 '1-2节 8:00-9:40'"""
    found = {}
    # 所有模式都要求时间在 06:00-23:00 合法范围内
    def ok_time(t):
        try:
            h, m = map(int, t.split(":"))
            return 6 <= h <= 23 and 0 <= m <= 59
        except: return False

    # 模式1: "1-2节 08:00-09:40" 或 "1-2节8:00-9:40"
    for m in re.finditer(r'(\d+-\d+)节\s*[:：]?\s*(\d{1,2}:\d{2})\s*[-~～]\s*(\d{1,2}:\d{2})', text):
        slot, start, end = m.group(1), m.group(2), m.group(3)
        if slot not in found and ok_time(start) and ok_time(end):
            found[slot] = (start, end)
    # 模式2: 表格行 "1-2 | 8:00 | 9:40"
    for m in re.finditer(r'(\d+-\d+)\s+(\d{1,2}:\d{2})\s+(\d{1,2}:\d{2})', text):
        slot, start, end = m.group(1), m.group(2), m.group(3)
        if slot not in found and ok_time(start) and ok_time(end):
            found[slot] = (start, end)
    # 模式3: "第一大节 8:00-9:40" 映射到 1-2
    CN_NUM = {"一": "1", "二": "2", "三": "3", "四": "4", "五": "5", "六": "6"}
    period_map = {"1": "1-2", "2": "3-4", "3": "5-6", "4": "7-8", "5": "9-10"}
    for m in re.finditer(r'第([一二三四五六])大节\s*(\d{1,2}:\d{2})\s*[-~～]\s*(\d{1,2}:\d{2})', text):
        cn = m.group(1)
        if cn in CN_NUM:
            slot = period_map.get(CN_NUM[cn])
            if slot and slot not in found and ok_time(m.group(2)) and ok_time(m.group(3)):
                found[slot] = (m.group(2), m.group(3))
    return found


def parse_weeks(desc: str, total_weeks: int = 18) -> list:
    """解析 '1-3周(单),7-9周,11周' → [1,3,7,8,9,11]"""
    weeks = set()
    # 先标准化分隔符和括号
    desc = desc.replace("（", "(").replace("）", ")").replace("，", ",").replace("、", ",").replace(" ", "")
    # 匹配: 数字-数字周(单), 数字周, 数字-数字周单, 等
    for m in re.finditer(r'(\d+)(?:-(\d+))?\s*周\s*[\(（]?([单双])?[\)）]?', desc):
        a = int(m.group(1))
        b = int(m.group(2)) if m.group(2) else a
        mod = m.group(3)  # "单" or "双" or None
        for w in range(a, b + 1):
            if mod == "单" and w % 2 == 0: continue
            if mod == "双" and w % 2 == 1: continue
            if 1 <= w <= total_weeks:
                weeks.add(w)
    return sorted(weeks)


def generate_ics(courses: list, start_date: datetime.date, slots: dict, cal_name: str = "课表") -> str:
    """生成 ICS 日历字符串"""
    lines = [
        "BEGIN:VCALENDAR", "VERSION:2.0",
        "PRODID:-//Schedule2ICS//CN", "CALSCALE:GREGORIAN",
        "METHOD:PUBLISH",
        f"X-WR-CALNAME:{cal_name}",
        "X-WR-TIMEZONE:Asia/Shanghai",
        "BEGIN:VTIMEZONE", "TZID:Asia/Shanghai",
        "BEGIN:STANDARD", "TZOFFSETFROM:+0800", "TZOFFSETTO:+0800",
        "TZNAME:CST", "DTSTART:19700101T000000",
        "END:STANDARD", "END:VTIMEZONE",
    ]

    uid = 0
    for name, days, slot_key, week_desc, location in courses:
        if slot_key not in slots:
            print(f"  ⚠ 未知节次 {slot_key}，跳过 {name}")
            continue
        start_t, end_t = slots[slot_key]
        weeks = parse_weeks(week_desc)

        for w in weeks:
            for d in days:
                date = start_date + datetime.timedelta(weeks=w - 1, days=DAY_MAP[d])
                ds = date.strftime("%Y%m%d")
                uid += 1
                lines.extend([
                    "BEGIN:VEVENT",
                    f"DTSTART;TZID=Asia/Shanghai:{ds}T{start_t.replace(':', '')}00",
                    f"DTEND;TZID=Asia/Shanghai:{ds}T{end_t.replace(':', '')}00",
                    f"SUMMARY:{name}",
                    f"LOCATION:{location}",
                    f"DESCRIPTION:第{w}周 {DAY_NAMES[DAY_MAP[d]]} {slot_key}节\\n{week_desc}",
                    f"UID:s2ics-{uid}",
                    "END:VEVENT",
                ])

    lines.append("END:VCALENDAR")
    return "\r\n".join(lines)


# ====== 交互式 CLI ======
if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="课表图片 → iOS 日历 .ics")
    ap.add_argument("file", help="课表文件路径 (PNG/JPG/PDF/DOCX)")
    ap.add_argument("--name", default="课表", help="学生姓名")
    ap.add_argument("--year", default="", help="学年，如 2025-2026")
    ap.add_argument("--semester", default="", help="学期，如 1 或 2")
    ap.add_argument("--start", default="2026-02-23", help="学期第一天（周一）YYYY-MM-DD")
    ap.add_argument("--weeks", type=int, default=18, help="总周数")
    ap.add_argument("--out", help="输出 .ics 路径")
    ap.add_argument("--slots", help="自定义节次时间，格式: '1-2=08:30-10:00;3-4=10:15-11:45'")
    args = ap.parse_args()

    # ====== 加载时间槽 ======
    SLOTS_CONFIG = SCRIPT_DIR / "slots.json"
    slots = dict(DEFAULT_SLOTS)

    # 1) 从本地配置文件加载（跨次复用）
    if SLOTS_CONFIG.exists():
        try:
            saved = json.load(open(SLOTS_CONFIG))
            if saved:
                slots = saved
        except: pass

    # 2) --slots 参数覆盖
    if args.slots:
        for item in args.slots.split(";"):
            item = item.strip()
            if "=" in item:
                key, val = item.split("=", 1)
                times = val.split("-")
                if len(times) == 2:
                    slots[key.strip()] = (times[0].strip(), times[1].strip())

    ftype = Path(args.file).suffix.upper()
    print(f"📄 识别中 ({ftype})...")
    text = extract_text(args.file)
    # 过滤明显损坏的 OCR 输出（纯乱码）
    clean = text.strip()
    chinese_chars = len(re.findall(r'[一-鿿]', clean))
    total_chars = len(clean.replace('\n', '').replace(' ', ''))
    quality = chinese_chars / max(total_chars, 1) if total_chars > 0 else 0
    if clean and quality > 0.02:  # 至少 2% 是中文字符
        print(clean[:2000])
    elif clean:
        print("  (OCR 输出质量过低，进入纯手动模式)")
        text = ""
    else:
        print("  (未识别到文字，进入纯手动模式)")

    # 3) 自动捕捉学年学期
    if not args.year or not args.semester:
        auto_year, auto_sem = detect_semester(text)
        if auto_year and not args.year:
            args.year = auto_year
        if auto_sem and not args.semester:
            args.semester = auto_sem
    if not args.year:
        args.year = input("学年？如 2025-2026\n> ").strip()
    if not args.semester:
        args.semester = input("第几学期？(1 或 2)\n> ").strip()

    # 4) 自动捕捉时间槽
    detected = detect_slots(text)
    if detected:
        slots.update(detected)
        json.dump(slots, open(SLOTS_CONFIG, "w"), ensure_ascii=False, indent=2)

    # 5) 交互式确认/修改
    print(f"\n📅 当前节次时间配置：")
    for k in sorted(slots.keys(), key=lambda x: int(x.split("-")[0])):
        s, e = slots[k]
        tag = " 🆕(OCR自动捕捉)" if k in detected else ""
        print(f"   {k}节: {s} - {e}{tag}")
    if detected:
        print(f"\n✅ 已从图片中自动捕捉到 {len(detected)} 个节次时间")
    print()
    change = input("需要修改吗？输入新配置或回车确认\n格式: 1-2=08:30-10:00;3-4=10:15-11:45\n> ").strip()
    if change:
        for item in change.split(";"):
            item = item.strip()
            if "=" in item:
                key, val = item.split("=", 1)
                times = val.split("-")
                if len(times) == 2:
                    slots[key.strip()] = (times[0].strip(), times[1].strip())
        json.dump(slots, open(SLOTS_CONFIG, "w"), ensure_ascii=False, indent=2)
        print(f"✅ 已保存\n")
    elif not detected:
        print()

    print("---")
    print("👆 以上是 OCR 原文。请对照图片，按以下格式输入每门课：")
    print("   课程名, 星期(一二三四五), 节次(如3-4), 周次(如1-18周), 教室")
    print("   输入空行结束。示例: 大学英语,一二三,1-2,1-3周单7-9周,教B-302")
    print()

    courses = []
    while True:
        try:
            line = input("> ").strip()
        except (EOFError, KeyboardInterrupt):
            break
        if not line:
            break
        parts = [p.strip() for p in line.split(",")]
        if len(parts) < 4:
            print("  格式: 课程名,星期,节次,周次[,教室]")
            continue
        name, days_str, slot, week_desc = parts[:4]
        loc = parts[4] if len(parts) > 4 else ""
        days = [d for d in days_str if d in DAY_MAP]
        if not days:
            print("  ⚠ 星期格式错误，请用一二三四五")
            continue
        courses.append((name, days, slot, week_desc, loc))
        print(f"  ✓ {name} ({len(days)}天, {slot}节)")

    if not courses:
        print("未输入任何课程，退出。")
        sys.exit(0)

    start = datetime.date.fromisoformat(args.start)
    # 日历名称：学年+学期 或 回退到 name
    if args.year and args.semester:
        cal_name = f"{args.year}学年第{args.semester}学期课表"
    else:
        cal_name = args.name
    ics = generate_ics(courses, start, slots, cal_name)

    out_path = args.out or f"{Path.home()}/Downloads/{cal_name}.ics"
    with open(out_path, "w") as f:
        f.write(ics)

    print(f"\n✅ 已生成: {out_path}")
    print(f"   共 {ics.count('BEGIN:VEVENT')} 个事件")
    print(f"\n📱 导入 iOS：AirDrop 到 iPhone → 点开 → 添加到日历")
